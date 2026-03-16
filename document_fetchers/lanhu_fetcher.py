# lanhu_fetcher.py
import re
import json
import time
import sys
import warnings
import subprocess
from pathlib import Path
from typing import Optional, List, Dict, Tuple, Union
from dataclasses import dataclass

import fitz
import requests


@dataclass
class PDFPageContent:
    """PDF页面内容结构"""
    page_number: int
    text: str
    has_images: bool
    layout_type: str


class LanhuFetcher:
    """
    蓝湖PRD获取器 - 终极增强版（支持多文件批量上传）
    专门针对蓝湖导出的PDF优化（图片+文字混合）
    """

    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.9",
        })
        # 忽略警告
        warnings.filterwarnings("ignore")

    def fetch(self, source: str) -> str:
        """
        获取PRD内容 - 支持单文件或多文件（分号分隔）
        保持向后兼容，单文件时行为不变
        """
        source = source.strip()
        if not source:
            return ""

        # 检查是否是多文件（包含分号且包含文件路径特征）
        if self._is_multiple_files(source):
            return self.fetch_multiple(source)

        # 单文件处理（原有逻辑）
        return self._fetch_single(source)

    def _fetch_single(self, source: str) -> str:
        """处理单个文件（原有fetch逻辑）"""
        if self._is_lanhu_url(source):
            return self._fetch_from_url(source)
        elif self._is_local_file(source):
            return self._read_local_file(source)
        else:
            # 纯文本直接返回
            return source

    def fetch_multiple(self, sources_input: Union[str, List[str]]) -> str:
        """
        批量获取多个PRD文件内容并合并

        :param sources_input: 分号分隔的路径字符串 或 路径列表
        :return: 合并后的PRD内容
        """
        # 统一转换为列表
        if isinstance(sources_input, str):
            # 按分号分割，支持中英文分号
            paths = [p.strip() for p in re.split(r'[;；]', sources_input) if p.strip()]
        else:
            paths = sources_input

        if not paths:
            return ""

        if len(paths) == 1:
            # 只有一个文件，使用单文件处理
            return self._fetch_single(paths[0])

        # 批量处理多个文件
        print(f"检测到 {len(paths)} 个PRD文件，开始批量读取...")

        all_contents = []
        success_count = 0

        for i, path in enumerate(paths, 1):
            print(f"\n[{i}/{len(paths)}] 正在读取: {Path(path).name}")

            try:
                if self._is_lanhu_url(path):
                    content = self._fetch_from_url(path)
                elif self._is_local_file(path):
                    content = self._read_local_file(path)
                else:
                    # 当作纯文本
                    content = path

                if content and content.strip():
                    # 添加文件分隔标记
                    separator = f"\n\n{'=' * 60}\n"
                    separator += f"## PRD文档 {i}: {Path(path).name}\n"
                    separator += f"{'=' * 60}\n\n"

                    all_contents.append(separator + content.strip())
                    success_count += 1
                    print(f"    ✓ 成功读取 ({len(content)} 字符)")
                else:
                    print(f"    ⚠ 文件内容为空")

            except Exception as e:
                print(f"    ✗ 读取失败: {e}")
                # 继续处理其他文件，不中断
                continue

        # 合并所有内容
        if not all_contents:
            return "# 错误：所有PRD文件读取失败\n\n请检查文件路径是否正确，或文件是否损坏。"

        # 添加合并头信息
        header = f"""# 合并PRD文档 ({success_count}/{len(paths)} 个文件成功)
> 生成时间: {time.strftime('%Y-%m-%d %H:%M:%S')}
> 文件列表: {', '.join([Path(p).name for p in paths[:5]])}{'...' if len(paths) > 5 else ''}

"""

        combined = header + '\n'.join(all_contents)
        print(f"\n✅ 批量读取完成，共 {success_count}/{len(paths)} 个文件，总长度 {len(combined)} 字符")

        return combined

    def _is_multiple_files(self, text: str) -> bool:
        """判断是否是多文件输入（分号分隔且包含文件特征）"""
        # 必须包含分号
        if ';' not in text and '；' not in text:
            return False

        # 分割后检查每个部分是否像文件路径
        parts = [p.strip() for p in re.split(r'[;；]', text) if p.strip()]
        if len(parts) < 2:
            return False

        # 至少两个部分看起来像文件路径
        file_like_count = sum(1 for p in parts if self._looks_like_file_path(p))
        return file_like_count >= 2

    def _looks_like_file_path(self, text: str) -> bool:
        """判断文本是否像文件路径"""
        text = text.strip()

        # 检查是否是URL（排除）
        if text.startswith(('http://', 'https://')):
            return False

        # 检查文件扩展名
        has_extension = any(ext in text.lower() for ext in ['.pdf', '.docx', '.doc', '.md', '.txt', '.xmind'])

        # 检查路径特征
        path_indicators = [
            text.startswith(('/', './', '../', '~/')),  # Unix路径
            text.startswith(('C:\\', 'D:\\', 'E:\\', 'F:\\')),  # Windows路径
            '/' in text or '\\' in text,  # 包含路径分隔符
        ]

        return has_extension or any(path_indicators)

    def _is_lanhu_url(self, text: str) -> bool:
        """判断是否是蓝湖URL"""
        patterns = [r'lanhuapp\.com', r'lanhu\.com', r'/url/', r'/web/#/item/']
        return any(re.search(p, text, re.IGNORECASE) for p in patterns)

    def _is_local_file(self, text: str) -> bool:
        """判断是否是本地文件路径（支持单个路径）"""
        text = text.strip()

        # 如果是多文件格式，返回False（让fetch_multiple处理）
        if ';' in text or '；' in text:
            return False

        if text.startswith(('/', './', '../', '~/', 'C:\\', 'D:\\', 'E:\\')):
            return True
        if Path(text).expanduser().exists():
            return True
        if any(ext in text.lower() for ext in ['.md', '.txt', '.docx', '.pdf', '.xmind']):
            if not text.startswith('http'):
                return True
        return False

    # ==================== 本地文件读取 ====================

    def _read_local_file(self, file_path: str) -> str:
        """读取本地文件"""
        path = Path(file_path).expanduser().resolve()

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        ext = path.suffix.lower()

        if ext == '.pdf':
            return self._read_pdf_comprehensive(path)
        elif ext in ['.md', '.txt', '.html']:
            return self._read_text_file(path)
        elif ext == '.docx':
            return self._read_docx(path)
        elif ext == '.xmind':
            # 新增：支持XMind格式
            return self._read_xmind(path)
        else:
            return self._read_text_file(path)

    def _read_xmind(self, xmind_path: Path) -> str:
        """读取XMind文件"""
        import zipfile

        try:
            with zipfile.ZipFile(xmind_path, 'r') as zf:
                # 尝试读取content.xml
                if 'content.xml' in zf.namelist():
                    content = zf.read('content.xml').decode('utf-8')
                    return self._parse_xmind_xml(content, xmind_path.name)
                elif 'content.json' in zf.namelist():
                    content = zf.read('content.json').decode('utf-8')
                    return self._parse_xmind_json(content, xmind_path.name)
                else:
                    raise ValueError("无法找到content.xml或content.json")
        except Exception as e:
            raise Exception(f"XMind解析失败: {e}")

    def _read_text_file(self, path: Path) -> str:
        """读取文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'ascii']
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        with open(path, 'rb') as f:
            return f.read().decode('utf-8', errors='ignore')

    def _read_docx(self, path: Path) -> str:
        """读取Word文档"""
        try:
            from docx import Document
            doc = Document(path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
            return '\n\n'.join(paragraphs)
        except ImportError:
            raise ImportError("读取Word需要安装: pip install python-docx")
        except Exception as e:
            raise Exception(f"读取Word文档失败: {e}")

    def _fetch_from_url(self, url: str) -> str:
        """从蓝湖URL获取（占位，实际实现需要蓝湖API）"""
        # 这里可以实现蓝湖API调用
        return f"# 蓝湖URL内容\n暂不支持直接解析蓝湖URL，请导出PDF后上传。\nURL: {url}"


    def _parse_xmind_xml(self, xml_content: str, filename: str) -> str:
        """解析XMind XML格式"""
        import xml.etree.ElementTree as ET

        try:
            root = ET.fromstring(xml_content)
        except ET.ParseError as e:
            raise ValueError(f"XML解析错误: {e}")

        ns = {'xmap': 'urn:xmind:xmap:xmlns:content:2.0'}
        all_text = [f"# XMind: {filename}\n"]

        # 查找所有sheet
        sheets = root.findall('.//xmap:sheet', ns) or root.findall('.//sheet')

        for sheet in sheets:
            sheet_title_elem = sheet.find('.//xmap:title', ns) or sheet.find('.//title')
            if sheet_title_elem is not None:
                sheet_title = ''.join(sheet_title_elem.itertext()).strip()
                if sheet_title:
                    all_text.append(f"## {sheet_title}")

            # 获取根主题
            root_topic = sheet.find('.//xmap:topic', ns) or sheet.find('.//topic')
            if root_topic is not None:
                self._extract_topics_xml(root_topic, all_text, level=0, ns=ns)

        return '\n'.join(all_text)

    def _extract_topics_xml(self, topic_elem, lines: List[str], level: int, ns: dict):
        """递归提取XML主题"""
        title_elem = topic_elem.find('.//xmap:title', ns) or topic_elem.find('.//title')
        if title_elem is not None:
            title = ''.join(title_elem.itertext()).strip()
            if title:
                indent = "  " * level
                prefix = "- " if level > 0 else "# "
                lines.append(f"{indent}{prefix}{title}")

        # 处理子主题
        children = topic_elem.find('.//xmap:children', ns) or topic_elem.find('.//children')
        if children is not None:
            topics_groups = children.findall('.//xmap:topics', ns) or children.findall('.//topics')
            for topics_group in topics_groups:
                sub_topics = topics_group.findall('.//xmap:topic', ns) or topics_group.findall('.//topic')
                for sub_topic in sub_topics:
                    self._extract_topics_xml(sub_topic, lines, level + 1, ns)

    def _parse_xmind_json(self, json_content: str, filename: str) -> str:
        """解析XMind Zen JSON格式"""
        import json

        try:
            data = json.loads(json_content)
        except json.JSONDecodeError as e:
            raise ValueError(f"JSON解析错误: {e}")

        all_text = [f"# XMind: {filename}\n"]
        root_topic = data.get('rootTopic')

        if root_topic:
            self._extract_topics_json(root_topic, all_text, level=0)

        return '\n'.join(all_text)

    def _extract_topics_json(self, topic: dict, lines: List[str], level: int):
        """递归提取JSON主题"""
        title = topic.get('title', '')
        if title:
            indent = "  " * level
            prefix = "- " if level > 0 else "# "
            lines.append(f"{indent}{prefix}{title}")

        children = topic.get('children', {})
        attached = children.get('attached', [])
        for child in attached:
            self._extract_topics_json(child, lines, level + 1)

    # ... 保留其他原有方法不变（_read_pdf_comprehensive, _read_with_pdfplumber等）...
    # 为节省篇幅，这里省略，实际使用时保留原有所有方法

    def _read_pdf_comprehensive(self, pdf_path: Path) -> str:
        """综合PDF读取策略 - 原有方法"""
        print(f"正在读取PDF: {pdf_path.name} ({pdf_path.stat().st_size / 1024:.1f} KB)")

        results = []

        # 策略1: pdfplumber（提取文本层和表格）
        try:
            print("  尝试 pdfplumber...")
            text = self._read_with_pdfplumber(pdf_path)
            if len(text) > 200:
                results.append(("pdfplumber", text, len(text)))
                print(f"    ✓ 提取 {len(text)} 字符")
        except Exception as e:
            print(f"    ✗ 失败: {e}")

        # 策略2: PyMuPDF（更好的文字定位和OCR预备）
        try:
            print("  尝试 PyMuPDF...")
            text = self._read_with_pymupdf_advanced(pdf_path)
            if len(text) > 200:
                results.append(("pymupdf", text, len(text)))
                print(f"    ✓ 提取 {len(text)} 字符")
        except Exception as e:
            print(f"    ✗ 失败: {e}")

        # 策略3: OCR（如果上面都失败或内容太少，说明是图片PDF）
        if not results or max(r[2] for r in results) < 500:
            try:
                print("  尝试 OCR识别...")
                text = self._read_with_ocr(pdf_path)
                if len(text) > 100:
                    results.append(("ocr", text, len(text)))
                    print(f"    ✓ OCR识别 {len(text)} 字符")
            except Exception as e:
                print(f"    ✗ OCR失败: {e}")

        # 策略4: 转换为图片后OCR（最后手段）
        if not results:
            try:
                print("  尝试 PDF转图片+OCR...")
                text = self._pdf_to_images_ocr(pdf_path)
                if len(text) > 100:
                    results.append(("pdf2image+ocr", text, len(text)))
                    print(f"    ✓ 识别 {len(text)} 字符")
            except Exception as e:
                print(f"    ✗ 失败: {e}")

        # 选择最佳结果或合并
        if not results:
            return self._create_error_content(pdf_path, "无法提取PDF内容，可能是扫描件或加密PDF")

        # 选择最长的结果
        best = max(results, key=lambda x: x[2])
        print(f"  最佳结果: {best[0]} ({best[2]} 字符)")

        # 后处理优化
        optimized = self._optimize_content(best[1], pdf_path.name)
        return optimized

    # 保留所有原有的PDF读取方法...
    def _read_with_pdfplumber(self, pdf_path: Path) -> str:
        """使用pdfplumber读取PDF - 原有方法"""
        import pdfplumber

        all_texts = []

        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_text = f"\n=== 第{i}页 ===\n"

                # 提取表格
                try:
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = self._format_table(table)
                            if table_text:
                                page_text += f"\n【表格】\n{table_text}\n"
                except:
                    pass

                # 提取文本（多种策略）
                text = ""

                # 策略A: 标准提取
                try:
                    text = page.extract_text() or ""
                except:
                    pass

                # 策略B: layout模式
                if len(text) < 50:
                    try:
                        text = page.extract_text(layout=True) or ""
                    except:
                        pass

                # 策略C: 宽松模式
                if len(text) < 50:
                    try:
                        text = page.extract_text(x_tolerance=5, y_tolerance=5) or ""
                    except:
                        pass

                if text.strip():
                    page_text += text.strip()

                # 检测图片
                try:
                    if page.images:
                        page_text += f"\n[本页包含 {len(page.images)} 个图片/原型图]"
                except:
                    pass

                all_texts.append(page_text)

        return '\n\n'.join(all_texts)

    def _read_with_pymupdf_advanced(self, pdf_path: Path) -> str:
        """使用PyMuPDF高级读取 - 原有方法"""
        import fitz

        doc = fitz.open(pdf_path)
        all_pages = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = f"\n=== 第{page_num + 1}页 ===\n"

            blocks = page.get_text("dict")["blocks"]
            blocks.sort(key=lambda b: b.get("bbox", [0, 0, 0, 0])[1])

            text_parts = []
            for block in blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        line_text = ""
                        for span in line.get("spans", []):
                            line_text += span.get("text", "")

                        if line.get("spans"):
                            size = line["spans"][0].get("size", 12)
                            flags = line["spans"][0].get("flags", 0)
                            if size > 14 or flags & 2 ** 4:
                                line_text = f"【{line_text}】"

                        if line_text.strip():
                            text_parts.append(line_text)

            if text_parts:
                page_text += '\n'.join(text_parts)

            images = page.get_images()
            drawings = page.get_drawings()

            if images or drawings:
                page_text += f"\n[包含 {len(images)} 图片, {len(drawings)} 图形元素]"

            all_pages.append(page_text)

        doc.close()
        return '\n\n'.join(all_pages)

    def _read_with_ocr(self, pdf_path: Path) -> str:
        """使用OCR读取PDF - 原有方法"""
        try:
            return self._ocr_with_tesseract(pdf_path)
        except Exception as e:
            print(f"    Tesseract OCR失败: {e}")

        try:
            return self._ocr_with_paddle(pdf_path)
        except Exception as e:
            print(f"    PaddleOCR失败: {e}")

        try:
            return self._ocr_with_easyocr(pdf_path)
        except Exception as e:
            print(f"    EasyOCR失败: {e}")

        raise Exception("所有OCR方案均失败")

    def _ocr_with_tesseract(self, pdf_path: Path) -> str:
        """Tesseract OCR - 原有方法"""
        from pdf2image import convert_from_path
        import pytesseract

        print("    使用Tesseract OCR...")
        images = convert_from_path(str(pdf_path), dpi=200)

        all_text = []
        for i, image in enumerate(images, 1):
            print(f"      识别第{i}/{len(images)}页...")
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            if text.strip():
                all_text.append(f"\n=== 第{i}页 ===\n{text.strip()}")

        return '\n\n'.join(all_text)

    def _ocr_with_paddle(self, pdf_path: Path) -> str:
        """PaddleOCR - 原有方法"""
        from pdf2image import convert_from_path
        from paddleocr import PaddleOCR

        print("    使用PaddleOCR...")
        ocr = PaddleOCR(use_angle_cls=True, lang='ch', show_log=False)
        images = convert_from_path(str(pdf_path), dpi=200)

        all_text = []
        for i, image in enumerate(images, 1):
            print(f"      识别第{i}/{len(images)}页...")
            temp_path = f"/tmp/ocr_temp_{i}.png"
            image.save(temp_path)
            result = ocr.ocr(temp_path, cls=True)

            if result and result[0]:
                texts = [line[1][0] for line in result[0]]
                page_text = '\n'.join(texts)
                all_text.append(f"\n=== 第{i}页 ===\n{page_text}")

            Path(temp_path).unlink(missing_ok=True)

        return '\n\n'.join(all_text)

    def _ocr_with_easyocr(self, pdf_path: Path) -> str:
        """EasyOCR - 原有方法"""
        from pdf2image import convert_from_path
        import easyocr

        print("    使用EasyOCR...")
        reader = easyocr.Reader(['ch_sim', 'en'])
        images = convert_from_path(str(pdf_path), dpi=200)

        all_text = []
        for i, image in enumerate(images, 1):
            print(f"      识别第{i}/{len(images)}页...")
            temp_path = f"/tmp/ocr_temp_{i}.png"
            image.save(temp_path)
            result = reader.readtext(temp_path, detail=0)
            if result:
                page_text = '\n'.join(result)
                all_text.append(f"\n=== 第{i}页 ===\n{page_text}")

            Path(temp_path).unlink(missing_ok=True)

        return '\n\n'.join(all_text)

    def _format_table(self, table: List[List]) -> str:
        """格式化表格 - 原有方法"""
        if not table:
            return ""

        cleaned = []
        for row in table:
            cleaned_row = [str(cell).strip() if cell else "" for cell in row]
            if any(cleaned_row):
                cleaned.append(cleaned_row)

        if not cleaned:
            return ""

        max_cols = max(len(row) for row in cleaned)
        lines = []
        header = cleaned[0] + [""] * (max_cols - len(cleaned[0]))
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * max_cols) + " |")

        for row in cleaned[1:]:
            row = row + [""] * (max_cols - len(row))
            lines.append("| " + " | ".join(row[:max_cols]) + " |")

        return '\n'.join(lines)

    def _optimize_content(self, content: str, filename: str) -> str:
        """优化内容结构 - 原有方法"""
        content = re.sub(r'\n{5,}', '\n\n\n', content)

        header = f"""# PRD文档: {filename}

> 解析时间: {time.strftime('%Y-%m-%d %H:%M:%S')}
> 原始格式: PDF（蓝湖导出）
> 内容长度: {len(content)} 字符
> 注意: 蓝湖PDF可能包含大量原型图，文字识别可能不完整，建议对照原始设计稿核对

---

"""
        return header + content

    def _create_error_content(self, pdf_path: Path, error_msg: str) -> str:
        """创建错误提示内容 - 原有方法"""
        return f"""# PRD读取失败

## 错误信息
{error_msg}

## 文件信息
- 路径: {pdf_path}
- 大小: {pdf_path.stat().st_size / 1024:.1f} KB

## 建议解决方案

### 方案1: 使用蓝湖"导出Word"功能（推荐）
1. 在蓝湖页面点击「导出」→「导出Word」
2. 保存.docx文件
3. 使用本工具选择该文件

### 方案2: 复制粘贴文本
1. 在蓝湖页面选中PRD文本
2. 复制粘贴到工具的文本框

### 方案3: 安装OCR依赖（自动识别图片文字）
```bash
pip install pdf2image pytesseract paddleocr easyocr
# macOS额外需要:
brew install tesseract tesseract-lang
brew install poppler  # for pdf2image
"""
