import os
import re
from pathlib import Path
from typing import Union, Optional


class LocalDocumentFetcher:
    """本地文档读取器 - 支持多种格式"""

    SUPPORTED_EXTENSIONS = {
        '.md', '.markdown', '.txt', '.docx', '.pdf',
        '.doc', '.rtf', '.html', '.htm', '.json'
    }

    def __init__(self):
        self.content = ""
        self.file_type = ""

    def fetch(self, file_path: Union[str, Path]) -> str:
        """
        读取本地文档
        :param file_path: 文件路径或字符串
        :return: 文档文本内容
        """
        path = Path(file_path).expanduser().resolve()

        # 检查是否是字符串内容（用户直接粘贴的文本）
        if self._is_text_content(str(file_path)):
            return str(file_path)

        # 检查文件是否存在
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        if not path.is_file():
            raise ValueError(f"路径不是文件: {path}")

        # 检查扩展名
        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            # 尝试作为文本读取
            try:
                return self._read_text(path)
            except:
                raise ValueError(f"不支持的文件格式: {ext}，支持的格式: {self.SUPPORTED_EXTENSIONS}")

        self.file_type = ext

        # 根据格式选择读取方式
        if ext in ['.md', '.markdown', '.txt', '.json', '.html', '.htm', '.rtf']:
            return self._read_text(path)
        elif ext == '.docx':
            return self._read_docx(path)
        elif ext == '.doc':
            return self._read_doc(path)
        elif ext == '.pdf':
            return self._read_pdf(path)

        return self._read_text(path)

    def _is_text_content(self, content: str) -> bool:
        """判断是否是直接粘贴的文本内容（而非文件路径）"""
        # 如果包含换行符且长度较大，或包含常见Markdown标记，认为是文本内容
        if '\n' in content and len(content) > 100:
            # 检查是否包含Markdown特征
            md_patterns = ['# ', '## ', '- ', '* ', '> ', '```', '| ']
            if any(pattern in content for pattern in md_patterns):
                return True

        # 如果包含多个路径分隔符，认为是路径
        if content.count('/') > 2 or content.count('\\') > 2:
            return False

        # 如果以常见路径开头，认为是路径
        path_starts = ['/', './', '../', '~/', 'C:\\', 'D:\\', 'E:\\']
        if any(content.startswith(start) for start in path_starts):
            return False

        # 如果包含文件扩展名，认为是路径
        if any(ext in content for ext in ['.md', '.txt', '.docx', '.pdf']):
            return False

        return False

    def _read_text(self, path: Path) -> str:
        """读取文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'ascii']

        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        # 如果都失败，使用二进制读取然后忽略错误
        with open(path, 'rb') as f:
            return f.read().decode('utf-8', errors='ignore')

    def _read_docx(self, path: Path) -> str:
        """读取Word文档（.docx）"""
        try:
            from docx import Document
            doc = Document(path)

            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # 读取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(' | '.join(row_text))

            return '\n\n'.join(paragraphs)
        except ImportError:
            raise ImportError("读取Word需要安装: pip install python-docx")
        except Exception as e:
            raise Exception(f"读取Word文档失败: {e}")

    def _read_doc(self, path: Path) -> str:
        """读取旧版Word文档（.doc）"""
        # 尝试转换为docx或使用antiword
        try:
            # 方法1：使用antiword（macOS/Linux）
            import subprocess
            result = subprocess.run(['antiword', str(path)],
                                    capture_output=True, text=True)
            if result.returncode == 0:
                return result.stdout
        except:
            pass

        try:
            # 方法2：使用textract
            import textract
            return textract.process(str(path)).decode('utf-8')
        except ImportError:
            pass

        # 方法3：提示用户转换
        raise Exception(
            "不支持.doc格式，请转换为.docx后重试，或安装: pip install textract\n"
            "macOS转换命令: textutil -convert docx yourfile.doc"
        )

    def _read_pdf(self, path: Path) -> str:
        """读取PDF文档"""
        errors = []

        # 方法1：使用PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(path)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return '\n\n'.join(text_parts)
        except Exception as e:
            errors.append(f"PyPDF2: {e}")

        # 方法2：使用pdfplumber（效果更好）
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return '\n\n'.join(text_parts)
        except ImportError:
            errors.append("pdfplumber: 未安装")
        except Exception as e:
            errors.append(f"pdfplumber: {e}")

        # 方法3：使用pymupdf（fitz）
        try:
            import fitz  # pymupdf
            doc = fitz.open(path)
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            return '\n\n'.join(text_parts)
        except ImportError:
            errors.append("pymupdf: 未安装")
        except Exception as e:
            errors.append(f"pymupdf: {e}")

        raise Exception(f"PDF读取失败，尝试以下方法：\n" +
                        "\n".join(errors) +
                        "\n请安装: pip install pdfplumber 或 pip install PyMuPDF")

    def fetch_batch(self, file_paths: list) -> dict:
        """批量读取多个文件"""
        results = {}
        for path in file_paths:
            try:
                results[str(path)] = self.fetch(path)
            except Exception as e:
                results[str(path)] = f"ERROR: {e}"
        return results


# 便捷函数
def read_document(file_path: Union[str, Path]) -> str:
    """便捷函数：读取本地文档"""
    fetcher = LocalDocumentFetcher()
    return fetcher.fetch(file_path)


def is_valid_document(path: str) -> bool:
    """检查路径是否是有效的文档"""
    if not path or not isinstance(path, str):
        return False

    # 检查是否是文件路径
    p = Path(path).expanduser()
    if p.exists() and p.is_file():
        return True

    # 检查是否是文本内容（包含Markdown特征）
    if len(path) > 50 and ('\n' in path or '#' in path or '- ' in path):
        return True

    return False